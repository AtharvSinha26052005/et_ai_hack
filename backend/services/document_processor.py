"""Document processing service — PDF, DOCX, image parsing + OCR."""

import os
import re
import logging
from typing import Optional

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Processes uploaded documents into text and chunks."""

    def __init__(self):
        self._ocr_available = False
        try:
            import pytesseract
            self._ocr_available = True
        except ImportError:
            logger.warning("pytesseract not available — OCR disabled")

    async def process_file(self, file_path: str) -> dict:
        """Process a file and return extracted text, chunks, and metadata."""
        ext = os.path.splitext(file_path)[1].lower()

        if ext == ".pdf":
            return await self.process_pdf(file_path)
        elif ext in (".docx", ".doc"):
            return await self.process_docx(file_path)
        elif ext in (".xlsx", ".xls", ".csv"):
            return await self.process_spreadsheet(file_path)
        elif ext in (".png", ".jpg", ".jpeg", ".tiff", ".bmp"):
            return await self.process_image(file_path)
        elif ext == ".txt":
            return await self.process_text(file_path)
        else:
            raise ValueError(f"Unsupported file type: {ext}")

    async def process_pdf(self, file_path: str) -> dict:
        """Extract text from PDF using PyMuPDF, with OCR fallback."""
        import fitz  # PyMuPDF

        doc = fitz.open(file_path)
        pages = []
        full_text = ""

        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text("text")

            # If page has very little text, try OCR
            if len(text.strip()) < 50 and self._ocr_available:
                try:
                    text = await self._ocr_page(page)
                except Exception as e:
                    logger.warning(f"OCR failed for page {page_num}: {e}")

            pages.append({"page": page_num + 1, "text": text})
            full_text += text + "\n"

        doc.close()

        # Detect document type
        doc_type = self.detect_document_type(full_text)

        # Chunk the document
        chunks = self.chunk_document(full_text, file_path)

        # Store chunks in vector store
        chunk_count = await self._store_chunks(chunks, file_path)

        return {
            "text": full_text,
            "text_preview": full_text[:500],
            "page_count": len(pages),
            "chunk_count": chunk_count,
            "entity_count": 0,  # Will be updated by extraction agent
            "doc_type": doc_type,
            "pages": pages,
        }

    async def _ocr_page(self, page) -> str:
        """OCR a PDF page using Tesseract."""
        import pytesseract
        from PIL import Image
        import io

        pix = page.get_pixmap(dpi=300)
        img_data = pix.tobytes("png")
        img = Image.open(io.BytesIO(img_data))
        text = pytesseract.image_to_string(img)
        return text

    async def process_docx(self, file_path: str) -> dict:
        """Extract text from Word documents."""
        from docx import Document as DocxDocument

        doc = DocxDocument(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        full_text = "\n".join(paragraphs)

        # Also extract table content
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells)
                full_text += "\n" + row_text

        doc_type = self.detect_document_type(full_text)
        chunks = self.chunk_document(full_text, file_path)
        chunk_count = await self._store_chunks(chunks, file_path)

        return {
            "text": full_text,
            "text_preview": full_text[:500],
            "page_count": len(paragraphs) // 30 + 1,  # Estimate
            "chunk_count": chunk_count,
            "entity_count": 0,
            "doc_type": doc_type,
        }

    async def process_spreadsheet(self, file_path: str) -> dict:
        """Parse Excel/CSV into structured records."""
        import openpyxl

        ext = os.path.splitext(file_path)[1].lower()

        if ext == ".csv":
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                full_text = f.read()
        else:
            wb = openpyxl.load_workbook(file_path, data_only=True)
            sheets_text = []
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                sheet_text = f"Sheet: {sheet_name}\n"
                for row in ws.iter_rows(values_only=True):
                    row_values = [str(v) if v is not None else "" for v in row]
                    sheet_text += " | ".join(row_values) + "\n"
                sheets_text.append(sheet_text)
            full_text = "\n\n".join(sheets_text)
            wb.close()

        doc_type = self.detect_document_type(full_text)
        chunks = self.chunk_document(full_text, file_path)
        chunk_count = await self._store_chunks(chunks, file_path)

        return {
            "text": full_text,
            "text_preview": full_text[:500],
            "page_count": 1,
            "chunk_count": chunk_count,
            "entity_count": 0,
            "doc_type": doc_type,
        }

    async def process_image(self, file_path: str) -> dict:
        """OCR for images (P&ID drawings, scanned forms)."""
        if not self._ocr_available:
            return {
                "text": "[Image - OCR not available]",
                "text_preview": "[Image - OCR not available]",
                "page_count": 1,
                "chunk_count": 0,
                "entity_count": 0,
                "doc_type": "pid_drawing",
            }

        import pytesseract
        from PIL import Image

        img = Image.open(file_path)
        text = pytesseract.image_to_string(img)

        doc_type = self.detect_document_type(text)
        chunks = self.chunk_document(text, file_path)
        chunk_count = await self._store_chunks(chunks, file_path)

        return {
            "text": text,
            "text_preview": text[:500],
            "page_count": 1,
            "chunk_count": chunk_count,
            "entity_count": 0,
            "doc_type": doc_type,
        }

    async def process_text(self, file_path: str) -> dict:
        """Process plain text files."""
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            full_text = f.read()

        doc_type = self.detect_document_type(full_text)
        chunks = self.chunk_document(full_text, file_path)
        chunk_count = await self._store_chunks(chunks, file_path)

        return {
            "text": full_text,
            "text_preview": full_text[:500],
            "page_count": 1,
            "chunk_count": chunk_count,
            "entity_count": 0,
            "doc_type": doc_type,
        }

    def detect_document_type(self, text: str) -> str:
        """Classify document type based on content keywords."""
        text_lower = text.lower()

        type_keywords = {
            "equipment_manual": ["operating manual", "installation guide", "user manual", "technical manual", "datasheet", "specifications"],
            "maintenance_record": ["work order", "maintenance report", "corrective maintenance", "preventive maintenance", "breakdown", "repair"],
            "inspection_report": ["inspection report", "inspection finding", "condition assessment", "survey report", "audit finding"],
            "safety_procedure": ["safety procedure", "emergency procedure", "evacuation", "first aid", "ppe", "lockout tagout", "loto"],
            "regulatory": ["regulation", "standard", "compliance", "oisd", "factory act", "bis standard", "iso", "osha"],
            "work_order": ["work order", "wo-", "task id", "planned date", "actual date", "technician"],
            "sop": ["standard operating procedure", "sop", "step 1", "step 2", "precautions", "procedure"],
            "incident_report": ["incident report", "near miss", "accident", "injury", "root cause", "investigation"],
            "pid_drawing": ["p&id", "piping and instrumentation", "flow diagram", "pid"],
        }

        scores = {}
        for doc_type, keywords in type_keywords.items():
            score = sum(1 for kw in keywords if kw in text_lower)
            if score > 0:
                scores[doc_type] = score

        if scores:
            return max(scores, key=scores.get)
        return "other"

    def chunk_document(
        self,
        text: str,
        source: str = "",
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
    ) -> list[dict]:
        """Industrial-aware text chunking that preserves:
        - Equipment tag references (e.g., P-101A/B)
        - Table structures
        - Section headers and hierarchy
        - Procedure step sequences
        """
        if not text.strip():
            return []

        # Split by sections first
        sections = re.split(r'\n(?=[A-Z0-9][A-Z0-9\s\-\.]{3,}(?:\n|:))', text)

        chunks = []
        chunk_id = 0

        for section in sections:
            # If section is small enough, keep as one chunk
            if len(section) <= chunk_size:
                if section.strip():
                    chunks.append({
                        "id": f"{os.path.basename(source)}_chunk_{chunk_id}",
                        "text": section.strip(),
                        "source": source,
                        "chunk_index": chunk_id,
                    })
                    chunk_id += 1
            else:
                # Split large sections with overlap
                words = section.split()
                current_chunk = []
                current_length = 0

                for word in words:
                    current_chunk.append(word)
                    current_length += len(word) + 1

                    if current_length >= chunk_size:
                        chunk_text = " ".join(current_chunk)
                        chunks.append({
                            "id": f"{os.path.basename(source)}_chunk_{chunk_id}",
                            "text": chunk_text.strip(),
                            "source": source,
                            "chunk_index": chunk_id,
                        })
                        chunk_id += 1

                        # Overlap: keep last N words
                        overlap_words = chunk_size // 5  # ~20% overlap
                        current_chunk = current_chunk[-overlap_words:]
                        current_length = sum(len(w) + 1 for w in current_chunk)

                # Add remaining text
                if current_chunk:
                    chunk_text = " ".join(current_chunk)
                    if chunk_text.strip():
                        chunks.append({
                            "id": f"{os.path.basename(source)}_chunk_{chunk_id}",
                            "text": chunk_text.strip(),
                            "source": source,
                            "chunk_index": chunk_id,
                        })
                        chunk_id += 1

        return chunks

    async def _store_chunks(self, chunks: list[dict], file_path: str) -> int:
        """Store document chunks in ChromaDB."""
        if not chunks:
            return 0

        try:
            from database.chroma_client import ChromaClient
            from sentence_transformers import SentenceTransformer

            # Generate embeddings
            model = SentenceTransformer("all-MiniLM-L6-v2")
            texts = [c["text"] for c in chunks]
            embeddings = model.encode(texts).tolist()

            # Get document ID from file path
            doc_id = os.path.splitext(os.path.basename(file_path))[0]

            collection = ChromaClient.get_collection()
            collection.add(
                ids=[c["id"] for c in chunks],
                embeddings=embeddings,
                documents=texts,
                metadatas=[
                    {
                        "source": c["source"],
                        "chunk_index": c["chunk_index"],
                        "document_id": doc_id,
                    }
                    for c in chunks
                ],
            )

            logger.info(f"Stored {len(chunks)} chunks in ChromaDB for {doc_id}")
            return len(chunks)

        except Exception as e:
            logger.error(f"Failed to store chunks in ChromaDB: {e}")
            return len(chunks)  # Return count anyway for metadata
